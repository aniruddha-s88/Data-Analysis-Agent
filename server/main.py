from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import Response
from fastapi.middleware.cors import CORSMiddleware
from agent import ask_agent, ask_agent_with_trace, is_greeting
from history import append_message, get_history
from vector_db import get_vector_store
from database import DATABASE_URL
from cache import set_summary, get_summary
import pandas as pd
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.graphics.shapes import Drawing
from reportlab.graphics.charts.barcharts import VerticalBarChart
from reportlab.graphics.charts.lineplots import LinePlot
from reportlab.graphics.charts.piecharts import Pie
from datetime import datetime
import re

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

vector_db = get_vector_store()

@app.post("/upload-csv/")
async def upload_csv(file: UploadFile = File(...), chat_id: str | None = None):
    filename = (file.filename or "").lower()
    try:
        if filename.endswith(".xlsx") or filename.endswith(".xls"):
            df = pd.read_excel(file.file)
        else:
            df = pd.read_csv(file.file)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Failed to parse file: {exc}")
    table = _table_name_for_chat(chat_id)
    df.to_sql(table, DATABASE_URL, if_exists="replace")
    summary = _dataset_summary(df)
    set_summary(summary, chat_id=chat_id)
    return {"message": "CSV uploaded successfully"}

@app.post("/upload-pdf/")
async def upload_pdf(file: UploadFile = File(...), chat_id: str | None = None):
    text = file.file.read().decode("utf-8", errors="ignore")
    vector_db.add_texts([text])
    return {"message": "PDF uploaded"}

@app.get("/ask/")
def ask(query: str, chat_id: str | None = None):
    response, trace = ask_agent_with_trace(query, chat_id=chat_id)
    mode = "friendly" if is_greeting(query) else "business"
    append_message("user", query, mode)
    append_message("assistant", response, mode)
    return {"response": response, "trace": trace}


@app.get("/history/")
def history(limit: int = 50):
    return {"messages": get_history(limit=limit)}


@app.get("/report/")
def report(limit: int = 20, title: str = "Business AI Agent Report", chat_id: str | None = None):
    dataset_summary = get_summary(chat_id=chat_id)
    df = None
    try:
        df = pd.read_sql(f"SELECT * FROM {_table_name_for_chat(chat_id)}", DATABASE_URL)
    except Exception:
        df = None

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=50,
        rightMargin=50,
        topMargin=50,
        bottomMargin=50,
    )
    story = []
    styles = _report_styles()

    story.extend(_build_cover_page(title, styles))
    story.append(PageBreak())

    story.append(Paragraph("Executive Summary", styles["section"]))
    for line in _build_exec_summary(df, dataset_summary):
        story.append(Paragraph(f"• {line}", styles["body"]))
    story.append(Spacer(1, 16))

    story.append(Paragraph("Dataset Overview", styles["section"]))
    overview_lines = _build_dataset_overview(df, dataset_summary)
    for line in overview_lines:
        story.append(Paragraph(f"• {line}", styles["body"]))
    story.append(Spacer(1, 16))

    story.append(Paragraph("Key Metrics", styles["section"]))
    metrics_table = _build_metrics_table(df)
    if metrics_table:
        story.append(metrics_table)
        story.append(Spacer(1, 16))
    else:
        story.append(Paragraph("No numeric metrics available.", styles["body"]))

    chart = _build_chart(df)
    if chart:
        story.append(Paragraph("Performance Snapshot", styles["section"]))
        story.append(chart)
        story.append(Spacer(1, 16))

    line_chart = _build_line_chart(df)
    if line_chart:
        story.append(Paragraph("Trend Analysis", styles["section"]))
        story.append(line_chart)
        story.append(Spacer(1, 16))

    pie_charts = _build_pie_charts(df)
    if pie_charts:
        for title_text, pie in pie_charts:
            story.append(Paragraph(title_text, styles["section"]))
            story.append(pie)
            story.append(Spacer(1, 12))

    story.append(Paragraph("Recommendations", styles["section"]))
    for line in _build_recommendations(df):
        story.append(Paragraph(f"- {line}", styles["body"]))
    story.append(Spacer(1, 16))

    story.append(Paragraph("Sales Strategy", styles["section"]))
    for line in _build_sales_strategy(df):
        story.append(Paragraph(f"- {line}", styles["body"]))
    story.append(Spacer(1, 16))

    doc.build(story)
    buffer.seek(0)
    return Response(content=buffer.read(), media_type="application/pdf")


@app.get("/report-docx/")
def report_docx(limit: int = 20, title: str = "Business AI Agent Report", chat_id: str | None = None):
    try:
        from docx import Document
    except Exception:
        raise HTTPException(
            status_code=500,
            detail="Missing dependency: python-docx. Install with `pip install python-docx`.",
        )

    dataset_summary = get_summary(chat_id=chat_id)
    df = None
    try:
        df = pd.read_sql(f"SELECT * FROM {_table_name_for_chat(chat_id)}", DATABASE_URL)
    except Exception:
        df = None

    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Prepared by Business AI Agent • {datetime.now().strftime('%B %d, %Y')}")

    doc.add_heading("Executive Summary", level=1)
    for line in _build_exec_summary(df, dataset_summary):
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Dataset Overview", level=1)
    for line in _build_dataset_overview(df, dataset_summary):
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Key Metrics", level=1)
    metrics = _build_metrics_rows(df)
    if metrics:
        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Metric"
        hdr_cells[1].text = "Mean"
        hdr_cells[2].text = "Min"
        hdr_cells[3].text = "Max"
        hdr_cells[4].text = "Std"
        for row in metrics:
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
    else:
        doc.add_paragraph("No numeric metrics available.")

    doc.add_heading("Recommendations", level=1)
    for line in _build_recommendations(df):
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Sales Strategy", level=1)
    for line in _build_sales_strategy(df):
        doc.add_paragraph(line, style="List Bullet")

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return Response(
        content=output.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.get("/dashboard-data/")
def dashboard_data(chat_id: str | None = None):
    df = None
    try:
        df = pd.read_sql(f"SELECT * FROM {_table_name_for_chat(chat_id)}", DATABASE_URL)
    except Exception:
        df = None

    if df is None or df.empty:
        raise HTTPException(status_code=404, detail="No dataset available.")

    numeric_cols = df.select_dtypes(include="number")
    categorical_cols = df.select_dtypes(exclude="number")

    metric_col = _pick_metric_column(numeric_cols)
    category_col = _pick_category_column(categorical_cols, ["Category", "Segment", "Product", "State", "Region"])
    region_col = _pick_category_column(categorical_cols, ["Region", "State", "Country", "City", "Market"])

    category_breakdown = _aggregate_top(df, category_col, metric_col, "Category Mix")
    region_breakdown = _aggregate_top(df, region_col, metric_col, "Regional Performance")
    trend = _aggregate_trend(df, metric_col)

    kpis = _build_kpis(df, metric_col)

    return {
        "metric": metric_col or "Metric",
        "kpis": kpis,
        "category_breakdown": category_breakdown,
        "region_breakdown": region_breakdown,
        "trend": trend,
    }


def _report_styles():
    styles = getSampleStyleSheet()
    brand_primary = colors.HexColor("#0F172A")
    brand_accent = colors.HexColor("#0EA5E9")
    brand_muted = colors.HexColor("#64748B")
    brand_bg = colors.HexColor("#F8FAFC")
    styles.add(
        ParagraphStyle(
            name="cover_title",
            parent=styles["Title"],
            fontSize=28,
            leading=32,
            textColor=brand_primary,
        )
    )
    styles.add(
        ParagraphStyle(
            name="cover_subtitle",
            parent=styles["BodyText"],
            fontSize=12,
            leading=16,
            textColor=brand_muted,
        )
    )
    styles.add(
        ParagraphStyle(
            name="kicker",
            parent=styles["BodyText"],
            fontSize=10,
            leading=12,
            textColor=brand_accent,
            uppercase=True,
        )
    )
    styles.add(
        ParagraphStyle(
            name="section",
            parent=styles["Heading2"],
            textColor=brand_primary,
            spaceAfter=6,
            spaceBefore=12,
        )
    )
    styles.add(
        ParagraphStyle(
            name="body",
            parent=styles["BodyText"],
            fontSize=10,
            leading=14,
            textColor=brand_primary,
        )
    )
    styles.add(
        ParagraphStyle(
            name="centered",
            parent=styles["BodyText"],
            alignment=TA_CENTER,
            textColor=brand_muted,
        )
    )
    styles.add(
        ParagraphStyle(
            name="small_muted",
            parent=styles["BodyText"],
            fontSize=9,
            leading=12,
            textColor=brand_muted,
        )
    )
    styles.add(
        ParagraphStyle(
            name="chip",
            parent=styles["BodyText"],
            fontSize=9,
            leading=12,
            textColor=brand_primary,
            backColor=brand_bg,
            borderColor=brand_accent,
            borderWidth=0.6,
            borderPadding=(4, 6, 4, 6),
        )
    )
    return styles


def _build_cover_page(title, styles):
    subtitle = "Executive analytics brief"
    date_line = datetime.now().strftime("%B %d, %Y")
    story = [
        Paragraph("BUSINESS INTELLIGENCE REPORT", styles["kicker"]),
        Spacer(1, 10),
        Paragraph(title, styles["cover_title"]),
        Spacer(1, 8),
        Paragraph(subtitle, styles["cover_subtitle"]),
        Spacer(1, 18),
        Paragraph(f"Prepared on {date_line}", styles["centered"]),
        Spacer(1, 24),
        Paragraph("Prepared for: Leadership Team", styles["chip"]),
        Spacer(1, 220),
        Paragraph("Confidential • Internal Use Only", styles["small_muted"]),
    ]
    return story


def _build_metrics_table(df):
    rows = _build_metrics_rows(df)
    if not rows:
        return None
    data = [["Metric", "Mean", "Min", "Max", "Std Dev"]] + rows
    table = Table(data, colWidths=[150, 80, 80, 80, 80])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0F172A")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("ALIGN", (1, 1), (-1, -1), "RIGHT"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#F8FAFC"), colors.HexColor("#EEF2F6")]),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    return table


def _build_metrics_rows(df):
    if df is None or df.empty:
        return []
    numeric_cols = df.select_dtypes(include="number")
    if numeric_cols.empty:
        return []
    desc = numeric_cols.describe().transpose()
    rows = []
    for col, row in desc.iterrows():
        rows.append(
            [
                str(col),
                f"{row.get('mean', 0):.2f}",
                f"{row.get('min', 0):.2f}",
                f"{row.get('max', 0):.2f}",
                f"{row.get('std', 0):.2f}",
            ]
        )
        if len(rows) >= 8:
            break
    return rows


def _build_chart(df):
    if df is None or df.empty:
        return None

    numeric_cols = df.select_dtypes(include="number")
    if numeric_cols.empty:
        return None

    categorical_cols = df.select_dtypes(exclude="number")
    if categorical_cols.empty:
        return None

    preferred_cats = ["Category", "Region", "Segment", "State", "Country", "Customer", "Product"]
    cat_col = None
    for name in preferred_cats:
        if name in categorical_cols.columns:
            cat_col = name
            break
    if cat_col is None:
        cat_col = categorical_cols.columns[0]

    preferred_metrics = ["Sales", "Revenue", "Profit", "Amount"]
    metric_col = None
    for name in preferred_metrics:
        if name in numeric_cols.columns:
            metric_col = name
            break
    if metric_col is None:
        metric_col = numeric_cols.columns[0]

    grouped = (
        df.groupby(cat_col, dropna=False)[metric_col]
        .sum()
        .sort_values(ascending=False)
        .head(8)
    )
    if grouped.empty:
        return None

    drawing = Drawing(440, 220)
    chart = VerticalBarChart()
    chart.x = 40
    chart.y = 30
    chart.height = 160
    chart.width = 360
    chart.data = [list(grouped.values)]
    chart.categoryAxis.categoryNames = [str(v)[:12] for v in grouped.index]
    chart.valueAxis.valueMin = 0
    chart.bars[0].fillColor = colors.HexColor("#0EA5E9")
    drawing.add(chart)
    return drawing


def _build_line_chart(df):
    if df is None or df.empty:
        return None

    numeric_cols = df.select_dtypes(include="number")
    if numeric_cols.empty:
        return None

    metric_col = _pick_metric_column(numeric_cols)
    if not metric_col:
        return None

    date_col = _pick_date_column(df)
    if not date_col:
        return None

    work = df[[date_col, metric_col]].copy()
    work[date_col] = pd.to_datetime(work[date_col], errors="coerce")
    work = work.dropna(subset=[date_col])
    if work.empty:
        return None

    work["period"] = work[date_col].dt.to_period("M").astype(str)
    grouped = work.groupby("period")[metric_col].sum().sort_index().tail(12)
    if grouped.empty:
        return None

    points = list(enumerate(grouped.values))
    drawing = Drawing(440, 220)
    chart = LinePlot()
    chart.x = 40
    chart.y = 30
    chart.height = 160
    chart.width = 360
    chart.data = [points]
    chart.lines[0].strokeColor = colors.HexColor("#0EA5E9")
    chart.xValueAxis.valueMin = 0
    chart.xValueAxis.valueMax = max(0, len(points) - 1)
    chart.yValueAxis.valueMin = 0
    chart.yValueAxis.valueMax = max(grouped.values) * 1.1
    drawing.add(chart)
    return drawing


def _build_pie_charts(df):
    if df is None or df.empty:
        return []

    numeric_cols = df.select_dtypes(include="number")
    if numeric_cols.empty:
        return []

    categorical_cols = df.select_dtypes(exclude="number")
    if categorical_cols.empty:
        return []

    preferred_metrics = ["Sales", "Revenue", "Profit", "Amount"]
    metric_col = None
    for name in preferred_metrics:
        if name in numeric_cols.columns:
            metric_col = name
            break
    if metric_col is None:
        metric_col = numeric_cols.columns[0]

    chart_specs = []
    cat_col = _pick_category_column(categorical_cols, ["Category", "Segment", "Product"])
    if cat_col:
        chart_specs.append(("Category Mix", cat_col))
    region_col = _pick_category_column(categorical_cols, ["Region", "State", "Country"])
    if region_col and region_col != cat_col:
        chart_specs.append(("Regional Mix", region_col))

    drawings = []
    for title_text, col in chart_specs:
        grouped = (
            df.groupby(col, dropna=False)[metric_col]
            .sum()
            .sort_values(ascending=False)
            .head(6)
        )
        if grouped.empty:
            continue
        drawing = Drawing(440, 220)
        pie = Pie()
        pie.x = 130
        pie.y = 30
        pie.width = 180
        pie.height = 180
        pie.data = list(grouped.values)
        pie.labels = [str(v)[:12] for v in grouped.index]
        palette = [
            colors.HexColor("#0EA5E9"),
            colors.HexColor("#22C55E"),
            colors.HexColor("#F97316"),
            colors.HexColor("#A855F7"),
            colors.HexColor("#FACC15"),
            colors.HexColor("#38BDF8"),
        ]
        for i, _ in enumerate(pie.data):
            pie.slices[i].fillColor = palette[i % len(palette)]
        drawing.add(pie)
        drawings.append((title_text, drawing))
    return drawings


def _build_exec_summary(df, dataset_summary):
    lines = []
    if df is not None and not df.empty:
        lines.append(
            f"Dataset loaded with {len(df)} rows and {len(df.columns)} columns."
        )
    elif dataset_summary:
        lines.append("Dataset summary available from last upload.")
    else:
        lines.append("No dataset detected. Upload a CSV or Excel file to enable analysis.")
    return lines


def _build_recommendations(df):
    if df is None or df.empty:
        return ["Upload a dataset to generate recommendations."]

    recs = []
    numeric_cols = df.select_dtypes(include="number")
    categorical_cols = df.select_dtypes(exclude="number")

    preferred_metrics = ["Sales", "Revenue", "Profit", "Amount"]
    metric_col = None
    for name in preferred_metrics:
        if name in numeric_cols.columns:
            metric_col = name
            break
    if metric_col is None and not numeric_cols.empty:
        metric_col = numeric_cols.columns[0]

    if metric_col and not categorical_cols.empty:
        cat_col = _pick_category_column(categorical_cols, ["Category", "Segment", "Region", "Product", "State"])
        grouped = (
            df.groupby(cat_col, dropna=False)[metric_col]
            .sum()
            .sort_values(ascending=False)
        )
        if not grouped.empty:
            top_label = str(grouped.index[0])
            recs.append(f"Double down on top-performing {cat_col} '{top_label}' to maximize {metric_col}.")
        if len(grouped) > 3:
            bottom_label = str(grouped.index[-1])
            recs.append(f"Investigate underperforming {cat_col} '{bottom_label}' for pricing or promotion gaps.")

    date_col = _pick_date_column(df)
    if metric_col and date_col:
        work = df[[date_col, metric_col]].copy()
        work[date_col] = pd.to_datetime(work[date_col], errors="coerce")
        work = work.dropna(subset=[date_col])
        if not work.empty:
            work["period"] = work[date_col].dt.to_period("M").astype(str)
            trend = work.groupby("period")[metric_col].sum().sort_index()
            if len(trend) >= 6:
                recent = trend.tail(3).mean()
                prior = trend.tail(6).head(3).mean()
                if prior and recent < prior:
                    recs.append(f"Recent {metric_col} is trending down; consider targeted campaigns in the last 3 months.")
                elif prior and recent > prior:
                    recs.append(f"Recent {metric_col} is trending up; scale campaigns that performed well in recent months.")

    nulls = df.isna().sum().sum()
    if nulls > 0:
        recs.append("Improve data quality by filling missing values to enable more accurate insights.")

    if not recs:
        recs.append("Add more granular dimensions (region, product, channel) to unlock deeper recommendations.")
    return recs


def _build_sales_strategy(df):
    if df is None or df.empty:
        return ["Upload a dataset to generate a tailored sales strategy."]

    strategy = []
    numeric_cols = df.select_dtypes(include="number")
    categorical_cols = df.select_dtypes(exclude="number")
    metric_col = _pick_metric_column(numeric_cols)

    if metric_col and not categorical_cols.empty:
        channel_col = _pick_category_column(categorical_cols, ["Channel", "Segment", "Category", "Region"])
        if channel_col:
            grouped = (
                df.groupby(channel_col, dropna=False)[metric_col]
                .sum()
                .sort_values(ascending=False)
                .head(3)
            )
            if not grouped.empty:
                top = ", ".join(str(v) for v in grouped.index)
                strategy.append(f"Focus promotions on top segments: {top} to lift {metric_col}.")

    date_col = _pick_date_column(df)
    if metric_col and date_col:
        work = df[[date_col, metric_col]].copy()
        work[date_col] = pd.to_datetime(work[date_col], errors="coerce")
        work = work.dropna(subset=[date_col])
        if not work.empty:
            work["period"] = work[date_col].dt.to_period("M").astype(str)
            trend = work.groupby("period")[metric_col].sum().sort_index().tail(6)
            if len(trend) >= 4:
                latest = trend.tail(2).mean()
                earlier = trend.head(2).mean()
                if latest < earlier:
                    strategy.append("Recent performance softened; run limited-time offers and re-activate dormant customers.")
                else:
                    strategy.append("Momentum is positive; increase inventory in top sellers and scale winning channels.")

    if not strategy:
        strategy.append("Add customer, channel, and campaign fields to unlock stronger sales strategy insights.")
    return strategy

def _build_dataset_overview(df, dataset_summary):
    lines = []
    if df is None or df.empty:
        if dataset_summary:
            lines.append("Cached dataset overview:")
            lines.extend(dataset_summary.splitlines()[:12])
        else:
            lines.append("No dataset overview available.")
        return lines

    numeric_cols = df.select_dtypes(include="number")
    lines.append(f"Columns: {', '.join(df.columns[:12])}" + ("..." if len(df.columns) > 12 else ""))
    if not numeric_cols.empty:
        lines.append(f"Numeric columns: {', '.join(numeric_cols.columns[:10])}" + ("..." if len(numeric_cols.columns) > 10 else ""))
    nulls = df.isna().sum().sort_values(ascending=False)
    top_nulls = nulls[nulls > 0].head(5)
    if not top_nulls.empty:
        lines.append("Top missing fields:")
        for col, count in top_nulls.items():
            lines.append(f"- {col}: {int(count)} missing")
    else:
        lines.append("No missing values detected.")
    return lines


def _pick_metric_column(numeric_cols):
    if numeric_cols is None or numeric_cols.empty:
        return None
    preferred = ["Sales", "Revenue", "Profit", "Amount", "Total"]
    for name in preferred:
        if name in numeric_cols.columns:
            return name
    return numeric_cols.columns[0]


def _pick_category_column(categorical_cols, preferred_names):
    if categorical_cols is None or categorical_cols.empty:
        return None
    for name in preferred_names:
        if name in categorical_cols.columns:
            return name
    return categorical_cols.columns[0]


def _aggregate_top(df, category_col, metric_col, title):
    if not category_col or not metric_col:
        return {"title": title, "labels": [], "values": [], "metric_label": "Metric"}
    grouped = (
        df.groupby(category_col, dropna=False)[metric_col]
        .sum()
        .sort_values(ascending=False)
        .head(6)
    )
    return {
        "title": title,
        "labels": [str(v) for v in grouped.index],
        "values": [float(v) for v in grouped.values],
        "metric_label": metric_col,
    }


def _aggregate_trend(df, metric_col):
    if not metric_col:
        return {"title": "Trend Over Time", "labels": [], "values": [], "metric_label": "Metric"}
    date_col = _pick_date_column(df)
    if not date_col:
        return {"title": "Trend Over Time", "labels": [], "values": [], "metric_label": metric_col}

    work = df.copy()
    work[date_col] = pd.to_datetime(work[date_col], errors="coerce")
    work = work.dropna(subset=[date_col])
    if work.empty:
        return {"title": "Trend Over Time", "labels": [], "values": [], "metric_label": metric_col}

    work["period"] = work[date_col].dt.to_period("M").astype(str)
    grouped = work.groupby("period")[metric_col].sum().sort_index().tail(12)
    return {
        "title": "Trend Over Time",
        "labels": [str(v) for v in grouped.index],
        "values": [float(v) for v in grouped.values],
        "metric_label": metric_col,
    }


def _pick_date_column(df):
    for col in df.columns:
        if "date" in str(col).lower():
            return col
    return None


def _build_kpis(df, metric_col):
    total_rows = len(df)
    total_metric = None
    avg_metric = None
    if metric_col:
        total_metric = df[metric_col].sum()
        avg_metric = df[metric_col].mean()

    kpis = [
        {"label": "Rows", "value": f"{total_rows:,}"},
    ]
    if metric_col:
        kpis.append({"label": f"Total {metric_col}", "value": f"{total_metric:,.2f}"})
        kpis.append({"label": f"Avg {metric_col}", "value": f"{avg_metric:,.2f}"})
    return kpis


def _table_name_for_chat(chat_id):
    if not chat_id:
        return "sales"
    safe = re.sub(r"[^a-zA-Z0-9_]", "_", str(chat_id))
    return f"sales_{safe}"




def _dataset_summary(df):
    if df is None or df.empty:
        return "No rows found in the sales table."

    numeric_cols = df.select_dtypes(include="number")
    stats = ""
    if not numeric_cols.empty:
        stats = numeric_cols.describe().transpose().to_string()

    nulls = df.isna().sum().sort_values(ascending=False)
    top_nulls = nulls[nulls > 0].head(10).to_string()

    return (
        "Top rows:\n"
        f"{df.head(5).to_string()}\n\n"
        "Column summary (numeric):\n"
        f"{stats if stats else 'No numeric columns found.'}\n\n"
        "Missing values (top 10):\n"
        f"{top_nulls if top_nulls else 'No missing values detected.'}"
    )
